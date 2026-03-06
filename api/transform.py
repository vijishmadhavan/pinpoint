"""File transformation & creation endpoints — write, generate, convert, compress, download, run."""

from __future__ import annotations

import os
import shutil

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel

from api.helpers import _check_safe

router = APIRouter()


# --- Write & Create (Segment 15B) ---


class WriteFileRequest(BaseModel):
    path: str
    content: str
    append: bool = False


@router.post("/write-file")
def write_file_endpoint(req: WriteFileRequest) -> dict:
    """Create or write a text file."""
    path = os.path.abspath(req.path)
    _check_safe(path)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    mode = "a" if req.append else "w"
    with open(path, mode, encoding="utf-8") as f:
        f.write(req.content)
    return {"success": True, "path": path, "size": os.path.getsize(path), "append": req.append}


class GenerateExcelRequest(BaseModel):
    path: str
    data: list  # list of dicts or list of lists
    columns: list | None = None
    sheet_name: str = "Sheet1"


@router.post("/generate-excel")
def generate_excel_endpoint(req: GenerateExcelRequest) -> dict:
    """Create an Excel file from data."""
    import pandas as pd

    path = os.path.abspath(req.path)
    _check_safe(path)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    try:
        df = pd.DataFrame(req.data, columns=req.columns)
        df.to_excel(path, sheet_name=req.sheet_name, index=False)
        return {"success": True, "path": path, "rows": len(df), "columns": list(df.columns)}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to create Excel: {e}")


class GenerateChartRequest(BaseModel):
    data: dict  # {"labels": [...], "values": [...]} or {"x": [...], "y": [...]}
    chart_type: str = "bar"  # bar, line, pie, scatter, hist
    title: str = ""
    xlabel: str = ""
    ylabel: str = ""
    output_path: str | None = None


@router.post("/generate-chart")
def generate_chart_endpoint(req: GenerateChartRequest) -> dict:
    """Generate a chart image using matplotlib."""
    import matplotlib

    matplotlib.use("Agg")
    import tempfile

    import matplotlib.pyplot as plt

    output = (
        os.path.abspath(req.output_path)
        if req.output_path
        else os.path.join(tempfile.gettempdir(), "pinpoint_charts", f"chart_{int(__import__('time').time())}.png")
    )
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    try:
        fig, ax = plt.subplots(figsize=(10, 6))
        labels = req.data.get("labels") or req.data.get("x", [])
        values = req.data.get("values") or req.data.get("y", [])

        if req.chart_type == "bar":
            ax.bar(labels, values)
        elif req.chart_type == "line":
            ax.plot(labels, values, marker="o")
        elif req.chart_type == "pie":
            ax.pie(values, labels=labels, autopct="%1.1f%%")
        elif req.chart_type == "scatter":
            ax.scatter(labels, values)
        elif req.chart_type == "hist":
            ax.hist(values, bins="auto")
        else:
            raise HTTPException(
                status_code=400, detail=f"Unknown chart type: {req.chart_type}. Use: bar, line, pie, scatter, hist"
            )

        if req.title:
            ax.set_title(req.title)
        if req.xlabel:
            ax.set_xlabel(req.xlabel)
        if req.ylabel:
            ax.set_ylabel(req.ylabel)
        plt.tight_layout()
        fig.savefig(output, dpi=150)
        plt.close(fig)
        return {"success": True, "path": output}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Chart generation failed: {e}")


# --- PDF Tools (Segment 15B) ---


class MergePdfRequest(BaseModel):
    paths: list[str]
    output_path: str


@router.post("/merge-pdf")
def merge_pdf_endpoint(req: MergePdfRequest) -> dict:
    """Merge multiple PDFs into one."""
    import fitz  # PyMuPDF

    output = os.path.abspath(req.output_path)
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    merged = fitz.open()
    try:
        for p in req.paths:
            p = os.path.abspath(p)
            _check_safe(p)
            if not os.path.exists(p):
                raise HTTPException(status_code=404, detail=f"File not found: {p}")
            doc = fitz.open(p)
            try:
                merged.insert_pdf(doc)
            finally:
                doc.close()
        merged.save(output)
        total_pages = len(merged)
    finally:
        merged.close()
    return {"success": True, "path": output, "total_pages": total_pages, "files_merged": len(req.paths)}


class SplitPdfRequest(BaseModel):
    path: str
    pages: str  # "1-5", "3,7,10", "1-3,5,8-10"
    output_path: str


@router.post("/split-pdf")
def split_pdf_endpoint(req: SplitPdfRequest) -> dict:
    """Extract specific pages from a PDF."""
    import fitz

    path = os.path.abspath(req.path)
    _check_safe(path)
    output = os.path.abspath(req.output_path)
    _check_safe(output)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")
    os.makedirs(os.path.dirname(output), exist_ok=True)

    # Parse page specification
    page_nums = set()
    for part in req.pages.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            for i in range(int(start), int(end) + 1):
                page_nums.add(i)
        else:
            page_nums.add(int(part))

    doc = fitz.open(path)
    try:
        # Convert 1-based to 0-based
        zero_based = sorted(p - 1 for p in page_nums if 1 <= p <= len(doc))
        if not zero_based:
            raise HTTPException(status_code=400, detail=f"No valid pages. PDF has {len(doc)} pages.")

        new_doc = fitz.open()
        try:
            for page_no in zero_based:
                new_doc.insert_pdf(doc, from_page=page_no, to_page=page_no)
            new_doc.save(output)
            extracted = len(new_doc)
            source_pages = len(doc)
        finally:
            new_doc.close()
    finally:
        doc.close()
    return {"success": True, "path": output, "pages_extracted": extracted, "source_pages": source_pages}


class OrganizePdfRequest(BaseModel):
    path: str
    pages: list[int]  # ordered list: [3, 1, 2, 5] = reorder; [1, 1, 2] = duplicate page 1
    output_path: str


@router.post("/organize-pdf")
def organize_pdf_endpoint(req: OrganizePdfRequest) -> dict:
    """Reorder, duplicate, or remove PDF pages. Pages list defines exact output order (1-based)."""
    import fitz

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")
    output = os.path.abspath(req.output_path)
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    doc = fitz.open(path)
    try:
        total = len(doc)
        valid_pages = [p for p in req.pages if 1 <= p <= total]
        if not valid_pages:
            raise HTTPException(status_code=400, detail=f"No valid pages. PDF has {total} pages.")

        new_doc = fitz.open()
        try:
            for p in valid_pages:
                new_doc.insert_pdf(doc, from_page=p - 1, to_page=p - 1)
            new_doc.save(output)
            result_pages = len(new_doc)
        finally:
            new_doc.close()
    finally:
        doc.close()
    return {"success": True, "path": output, "output_pages": result_pages, "source_pages": total}


class PdfToImagesRequest(BaseModel):
    path: str
    pages: str | None = None  # "1,3,5" or "1-3" or None for all
    dpi: int = 150
    output_folder: str | None = None


@router.post("/pdf-to-images")
def pdf_to_images_endpoint(req: PdfToImagesRequest) -> dict:
    """Render PDF pages as images."""
    import fitz

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    doc = fitz.open(path)
    total = len(doc)

    # Parse pages
    if req.pages:
        page_nums = set()
        for part in req.pages.split(","):
            part = part.strip()
            if "-" in part:
                a, b = part.split("-", 1)
                page_nums.update(range(int(a), int(b) + 1))
            else:
                page_nums.add(int(part))
        zero_based = sorted(p - 1 for p in page_nums if 1 <= p <= total)
    else:
        _PAGE_CAP = 50  # prevent OOM on 500-page PDFs (specify pages param for more)
        zero_based = list(range(min(total, _PAGE_CAP)))

    if not zero_based:
        doc.close()
        raise HTTPException(status_code=400, detail=f"No valid pages. PDF has {total} pages.")

    out_folder = req.output_folder or os.path.join(
        os.path.dirname(path), os.path.splitext(os.path.basename(path))[0] + "_pages"
    )
    os.makedirs(out_folder, exist_ok=True)

    zoom = req.dpi / 72
    mat = fitz.Matrix(zoom, zoom)
    saved = []
    for page_no in zero_based:
        page = doc[page_no]
        pix = page.get_pixmap(matrix=mat)
        img_path = os.path.join(out_folder, f"page_{page_no + 1}.png")
        pix.save(img_path)
        saved.append(img_path)

    doc.close()
    return {"success": True, "images": saved, "count": len(saved), "output_folder": out_folder}


class ImagesToPdfRequest(BaseModel):
    paths: list
    output_path: str


@router.post("/images-to-pdf")
def images_to_pdf_endpoint(req: ImagesToPdfRequest) -> dict:
    """Combine images into a single PDF."""
    from PIL import Image

    if not req.paths:
        raise HTTPException(status_code=400, detail="No image paths provided")

    output = os.path.abspath(req.output_path)
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    images = []
    try:
        for p in req.paths:
            p = os.path.abspath(p)
            _check_safe(p)
            if not os.path.exists(p):
                raise HTTPException(status_code=404, detail=f"Image not found: {p}")
            img = Image.open(p)
            if img.mode == "RGBA":
                original = img
                img = img.convert("RGB")
                original.close()
            images.append(img)

        if not images:
            raise HTTPException(status_code=400, detail="No valid images loaded")

        images[0].save(output, "PDF", save_all=True, append_images=images[1:])
    finally:
        for img in images:
            img.close()

    return {"success": True, "path": output, "pages": len(images)}


# --- Image Tools (Segment 15B) ---


class ResizeImageRequest(BaseModel):
    path: str
    width: int | None = None
    height: int | None = None
    quality: int = 85
    output_path: str | None = None


@router.post("/resize-image")
def resize_image_endpoint(req: ResizeImageRequest) -> dict:
    """Resize or compress an image."""
    from PIL import Image

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    img = Image.open(path)
    orig_size = img.size

    if req.width and req.height:
        img = img.resize((req.width, req.height), Image.LANCZOS)
    elif req.width:
        ratio = req.width / img.width
        img = img.resize((req.width, int(img.height * ratio)), Image.LANCZOS)
    elif req.height:
        ratio = req.height / img.height
        img = img.resize((int(img.width * ratio), req.height), Image.LANCZOS)

    output = os.path.abspath(req.output_path) if req.output_path else path
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)
    if img.mode == "RGBA" and output.lower().endswith((".jpg", ".jpeg")):
        img = img.convert("RGB")
    img.save(output, quality=req.quality)
    return {
        "success": True,
        "path": output,
        "original_size": list(orig_size),
        "new_size": list(img.size),
        "file_size": os.path.getsize(output),
    }


class ConvertImageRequest(BaseModel):
    path: str
    format: str  # jpg, png, webp, bmp
    output_path: str | None = None
    quality: int = 90


@router.post("/convert-image")
def convert_image_endpoint(req: ConvertImageRequest) -> dict:
    """Convert image to a different format."""
    from PIL import Image

    try:
        from pillow_heif import register_heif_opener

        register_heif_opener()
    except ImportError:
        pass

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    fmt = req.format.lower().lstrip(".")
    ext_map = {"jpg": ".jpg", "jpeg": ".jpg", "png": ".png", "webp": ".webp", "bmp": ".bmp"}
    if fmt not in ext_map:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}. Use: jpg, png, webp, bmp")

    output = os.path.abspath(req.output_path) if req.output_path else os.path.splitext(path)[0] + ext_map[fmt]
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    img = Image.open(path)
    if img.mode == "RGBA" and fmt in ("jpg", "jpeg"):
        img = img.convert("RGB")
    img.save(output, quality=req.quality)
    return {"success": True, "path": output, "format": fmt, "size": os.path.getsize(output)}


class CropImageRequest(BaseModel):
    path: str
    x: int
    y: int
    width: int
    height: int
    output_path: str | None = None


@router.post("/crop-image")
def crop_image_endpoint(req: CropImageRequest) -> dict:
    """Crop an image to specified dimensions."""
    from PIL import Image

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    img = Image.open(path)
    try:
        box = (req.x, req.y, req.x + req.width, req.y + req.height)
        cropped = img.crop(box)
        try:
            output = os.path.abspath(req.output_path) if req.output_path else path
            _check_safe(output)
            os.makedirs(os.path.dirname(output), exist_ok=True)
            cropped.save(output)
            result = {"success": True, "path": output, "crop_box": list(box), "new_size": list(cropped.size)}
        finally:
            cropped.close()
    finally:
        img.close()
    return result


# --- Image Metadata (EXIF) ---


class ImageMetadataRequest(BaseModel):
    path: str | None = None
    folder: str | None = None


def _extract_exif(filepath: str) -> dict:
    """Extract EXIF metadata from a single image file."""
    from PIL import Image
    from PIL.ExifTags import IFD

    try:
        from pillow_heif import register_heif_opener

        register_heif_opener()
    except ImportError:
        pass

    img = Image.open(filepath)
    result = {
        "path": filepath,
        "dimensions": {"width": img.width, "height": img.height},
        "format": img.format or os.path.splitext(filepath)[1].lstrip(".").upper(),
    }

    exif_data = img.getexif()
    if not exif_data:
        img.close()
        result["exif"] = None
        return result

    def _rational(val: object) -> float | None:
        """Convert IFDRational or tuple to float."""
        if val is None:
            return None
        if hasattr(val, "numerator"):
            return float(val)
        if isinstance(val, tuple) and len(val) == 2:
            return val[0] / val[1] if val[1] else 0
        return float(val)

    exif = {}
    # DateTimeOriginal
    dt = exif_data.get(36867) or exif_data.get(306)  # DateTimeOriginal or DateTime
    if dt:
        exif["date_taken"] = str(dt).replace(":", "-", 2)  # 2025:02:20 → 2025-02-20

    # Camera
    make = (exif_data.get(271) or "").strip()
    model = (exif_data.get(272) or "").strip()
    if model:
        # Avoid "Canon Canon EOS 5D"
        exif["camera"] = model if make and model.startswith(make) else f"{make} {model}".strip()

    # Lens
    ifd_exif = exif_data.get_ifd(IFD.Exif)
    lens = ifd_exif.get(42036) if ifd_exif else None
    if lens:
        exif["lens"] = str(lens)

    # Focal length
    fl = ifd_exif.get(37386) if ifd_exif else None
    if fl is not None:
        v = _rational(fl)
        exif["focal_length"] = f"{v:.0f}mm" if v else None

    # Aperture
    fn = ifd_exif.get(33437) if ifd_exif else None
    if fn is not None:
        v = _rational(fn)
        exif["aperture"] = f"f/{v:.1f}" if v else None

    # Shutter speed
    et = ifd_exif.get(33434) if ifd_exif else None
    if et is not None:
        v = _rational(et)
        if v and v < 1:
            exif["shutter_speed"] = f"1/{int(round(1 / v))}"
        elif v:
            exif["shutter_speed"] = f"{v:.1f}s"

    # ISO
    iso = ifd_exif.get(34855) if ifd_exif else None
    if iso is not None:
        exif["iso"] = int(iso) if not isinstance(iso, tuple) else int(iso[0])

    # GPS
    try:
        gps_ifd = exif_data.get_ifd(IFD.GPSInfo)
        if gps_ifd:

            def _dms_to_dd(dms: tuple, ref: str) -> float:
                d, m, s = [_rational(x) for x in dms]
                dd = d + m / 60 + s / 3600
                return -dd if ref in ("S", "W") else dd

            lat_dms = gps_ifd.get(2)
            lat_ref = gps_ifd.get(1)
            lon_dms = gps_ifd.get(4)
            lon_ref = gps_ifd.get(3)
            if lat_dms and lon_dms:
                exif["gps"] = {
                    "lat": round(_dms_to_dd(lat_dms, lat_ref), 6),
                    "lon": round(_dms_to_dd(lon_dms, lon_ref), 6),
                }
    except Exception:
        pass

    # Orientation
    orient = exif_data.get(274)
    if orient:
        exif["orientation"] = int(orient)

    img.close()
    result["exif"] = exif if exif else None
    return result


IMAGE_EXTS_EXIF = {
    ".jpg",
    ".jpeg",
    ".png",
    ".tiff",
    ".tif",
    ".heic",
    ".heif",
    ".webp",
    ".bmp",
    ".dng",
    ".cr2",
    ".nef",
    ".arw",
}


@router.post("/image-metadata")
def image_metadata_endpoint(req: ImageMetadataRequest) -> dict:
    """Extract EXIF metadata from photos."""
    if req.path:
        path = os.path.abspath(req.path)
        _check_safe(path)
        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail=f"File not found: {path}")
        try:
            result = _extract_exif(path)
            result["_hint"] = "Metadata extracted. Answer the user's question about this image."
            return result
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Cannot read image: {e}")

    elif req.folder:
        folder = os.path.abspath(req.folder)
        _check_safe(folder)
        if not os.path.isdir(folder):
            raise HTTPException(status_code=404, detail=f"Folder not found: {folder}")

        results = {}
        dates = []
        cameras = set()
        has_gps = 0
        total = 0

        for fname in sorted(os.listdir(folder)):
            if total >= 200:
                break
            ext = os.path.splitext(fname)[1].lower()
            if ext not in IMAGE_EXTS_EXIF:
                continue
            fpath = os.path.join(folder, fname)
            if not os.path.isfile(fpath):
                continue
            try:
                meta = _extract_exif(fpath)
                # Strip full path for batch — just filename
                meta["path"] = fname
                results[fname] = meta
                total += 1
                if meta.get("exif"):
                    if meta["exif"].get("date_taken"):
                        dates.append(meta["exif"]["date_taken"][:10])
                    if meta["exif"].get("camera"):
                        cameras.add(meta["exif"]["camera"])
                    if meta["exif"].get("gps"):
                        has_gps += 1
            except Exception:
                continue

        if not results:
            return {"folder": folder, "images_processed": 0, "results": {}, "_hint": "No images found in folder."}

        summary = {
            "date_range": f"{min(dates)} to {max(dates)}" if dates else None,
            "cameras": sorted(cameras) if cameras else [],
            "has_gps": has_gps,
            "total": total,
        }
        return {
            "folder": folder,
            "images_processed": total,
            "results": results,
            "summary": summary,
            "_hint": f"Metadata for {total} images. Summary shows date range and cameras used.",
        }

    else:
        raise HTTPException(status_code=400, detail="Provide path (single image) or folder (batch).")


# --- Archive Tools (Segment 15B) ---


class CompressFilesRequest(BaseModel):
    paths: list[str]
    output_path: str


@router.post("/compress-files")
def compress_files_endpoint(req: CompressFilesRequest) -> dict:
    """Compress files into a zip archive."""
    import zipfile

    output = os.path.abspath(req.output_path)
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    added = 0
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in req.paths:
            p = os.path.abspath(p)
            _check_safe(p)
            if not os.path.exists(p):
                continue
            if os.path.isdir(p):
                for root, dirs, files in os.walk(p):
                    for f in files:
                        fp = os.path.join(root, f)
                        arcname = os.path.relpath(fp, os.path.dirname(p))
                        zf.write(fp, arcname)
                        added += 1
            else:
                zf.write(p, os.path.basename(p))
                added += 1

    return {"success": True, "path": output, "files_added": added, "archive_size": os.path.getsize(output)}


class ExtractArchiveRequest(BaseModel):
    path: str
    output_path: str | None = None


@router.post("/extract-archive")
def extract_archive_endpoint(req: ExtractArchiveRequest) -> dict:
    """Extract a zip archive."""
    import zipfile

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    output = os.path.abspath(req.output_path) if req.output_path else os.path.splitext(path)[0]
    _check_safe(output)

    try:
        _MAX_EXTRACT_SIZE = 2 * 1024 * 1024 * 1024  # 2GB uncompressed limit
        with zipfile.ZipFile(path, "r") as zf:
            # Security: check for path traversal, symlinks, and zip bomb
            output_abs = os.path.abspath(output)
            total_size = 0
            for info in zf.infolist():
                target = os.path.join(output_abs, info.filename)
                # normpath resolves .. without hitting disk (faster than abspath)
                if not os.path.normpath(target).startswith(output_abs):
                    raise HTTPException(status_code=400, detail=f"Unsafe path in archive: {info.filename}")
                # Block symlinks (external_attr >> 28 == 0xA for symlinks on Unix)
                if info.external_attr >> 28 == 0xA:
                    raise HTTPException(status_code=400, detail=f"Symlink in archive not allowed: {info.filename}")
                total_size += info.file_size
                if total_size > _MAX_EXTRACT_SIZE:
                    raise HTTPException(status_code=400, detail=f"Archive too large (>{_MAX_EXTRACT_SIZE // (1024**3)}GB uncompressed)")
            zf.extractall(output)
            return {"success": True, "path": output_abs, "files_extracted": len(zf.namelist())}
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Not a valid zip file")


# --- Download (Segment 15B) ---


class DownloadUrlRequest(BaseModel):
    url: str
    save_path: str | None = None


@router.post("/download-url")
def download_url_endpoint(req: DownloadUrlRequest) -> dict:
    """Download a file from a URL."""
    import urllib.error
    import urllib.request

    url = req.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL must start with http:// or https://")

    # Determine save path
    if req.save_path:
        save_path = os.path.abspath(req.save_path)
    else:
        from urllib.parse import urlparse

        parsed = urlparse(url)
        filename = os.path.basename(parsed.path) or "downloaded_file"
        save_dir = os.path.join(os.path.expanduser("~"), "Downloads", "Pinpoint")
        os.makedirs(save_dir, exist_ok=True)
        save_path = os.path.join(save_dir, filename)

    _check_safe(save_path)
    os.makedirs(os.path.dirname(save_path), exist_ok=True)

    try:
        req_obj = urllib.request.Request(url, headers={"User-Agent": "Pinpoint/1.0"})
        with urllib.request.urlopen(req_obj, timeout=60) as resp:
            with open(save_path, "wb") as f:
                shutil.copyfileobj(resp, f)  # stream in chunks, not buffer entire file in RAM
        return {
            "success": True,
            "path": save_path,
            "size": os.path.getsize(save_path),
            "url": url,
            "_hint": "Use index_file to make this file searchable, or send_file to share it.",
        }
    except urllib.error.URLError as e:
        raise HTTPException(status_code=400, detail=f"Download failed: {e}")


# --- Run Python (Segment 15D) ---

PYTHON_WORK_DIR = "/tmp/pinpoint_python"
os.makedirs(PYTHON_WORK_DIR, exist_ok=True)


class RunPythonRequest(BaseModel):
    code: str
    timeout: int = 30


@router.post("/run-python")
def run_python_endpoint(req: RunPythonRequest) -> dict:
    """Execute Python code and return stdout + created files."""
    import contextlib
    import io
    import signal

    timeout = min(req.timeout, 120)

    # Snapshot files before execution
    before = set()
    for root, dirs, files in os.walk(PYTHON_WORK_DIR):
        for f in files:
            before.add(os.path.join(root, f))

    # Pre-loaded namespace
    namespace = {
        "__builtins__": __builtins__,
        "WORK_DIR": PYTHON_WORK_DIR,
    }

    # Lazy imports inside namespace
    setup_code = f"""
import os, sys, json, math, re, pathlib, shutil, glob, hashlib, datetime, io, csv
WORK_DIR = "{PYTHON_WORK_DIR}"
os.chdir(WORK_DIR)
try:
    import numpy as np
except ImportError:
    pass
try:
    import pandas as pd
except ImportError:
    pass
try:
    from PIL import Image, ImageDraw, ImageFont, ImageFilter
except ImportError:
    pass
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
except ImportError:
    pass
"""

    full_code = setup_code + "\n" + req.code

    stdout_capture = io.StringIO()
    stderr_capture = io.StringIO()

    # Timeout handler
    def timeout_handler(signum: int, frame: object) -> None:
        raise TimeoutError(f"Code execution timed out after {timeout}s")

    try:
        old_handler = signal.signal(signal.SIGALRM, timeout_handler)
        signal.alarm(timeout)

        with contextlib.redirect_stdout(stdout_capture), contextlib.redirect_stderr(stderr_capture):
            exec(full_code, namespace)

        signal.alarm(0)
        signal.signal(signal.SIGALRM, old_handler)
    except TimeoutError as e:
        signal.alarm(0)
        return {"success": False, "error": str(e), "stdout": stdout_capture.getvalue()[:5000]}
    except Exception as e:
        signal.alarm(0)
        return {"success": False, "error": f"{type(e).__name__}: {e}", "stdout": stdout_capture.getvalue()[:5000]}

    # Find new/modified files
    after = set()
    for root, dirs, files in os.walk(PYTHON_WORK_DIR):
        for f in files:
            after.add(os.path.join(root, f))
    new_files = sorted(after - before)

    stdout_text = stdout_capture.getvalue()[:10000]
    stderr_text = stderr_capture.getvalue()[:2000]

    result = {"success": True, "stdout": stdout_text}
    if stderr_text:
        result["stderr"] = stderr_text
    if new_files:
        result["files_created"] = new_files[:50]  # cap to prevent huge JSON responses
        if len(new_files) > 50:
            result["files_created_total"] = len(new_files)
    return result


# --- PDF Tools (Segment 22I — iLovePDF-inspired) ---


class CompressPdfRequest(BaseModel):
    path: str
    output_path: str | None = None


@router.post("/compress-pdf")
def compress_pdf_endpoint(req: CompressPdfRequest) -> dict:
    """Compress a PDF by removing unused objects and deflating streams."""
    import fitz

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    original_size = os.path.getsize(path)
    output = os.path.abspath(req.output_path) if req.output_path else path
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    doc = fitz.open(path)
    try:
        doc.save(output, garbage=4, deflate=True, clean=True)
    finally:
        doc.close()

    new_size = os.path.getsize(output)
    reduction = round((1 - new_size / original_size) * 100, 1) if original_size > 0 else 0
    return {
        "success": True,
        "path": output,
        "original_size": original_size,
        "compressed_size": new_size,
        "reduction_percent": reduction,
        "_hint": f"Compressed {reduction}% — {original_size} → {new_size} bytes.",
    }


class AddPageNumbersRequest(BaseModel):
    path: str
    output_path: str | None = None
    position: str = "bottom-center"  # bottom-left, bottom-center, bottom-right
    start: int = 1
    format: str = "{n}"  # e.g. "Page {n} of {total}", "{n}"


@router.post("/add-page-numbers")
def add_page_numbers_endpoint(req: AddPageNumbersRequest) -> dict:
    """Add page numbers to every page of a PDF."""
    import fitz

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    output = os.path.abspath(req.output_path) if req.output_path else path
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    doc = fitz.open(path)
    try:
        total = len(doc)

        for i, page in enumerate(doc):
            num = req.start + i
            text = req.format.replace("{n}", str(num)).replace("{total}", str(total))
            rect = page.rect
            margin = 36  # 0.5 inch
            font_size = 10

            if "left" in req.position:
                x = margin
            elif "right" in req.position:
                x = rect.width - margin - font_size * len(text) * 0.4
            else:
                x = rect.width / 2 - font_size * len(text) * 0.2

            y = rect.height - margin
            page.insert_text((x, y), text, fontsize=font_size, color=(0.4, 0.4, 0.4))

        doc.save(output)
    finally:
        doc.close()
    return {"success": True, "path": output, "pages_numbered": total}


class PdfToWordRequest(BaseModel):
    path: str
    output_path: str | None = None


@router.post("/pdf-to-word")
def pdf_to_word_endpoint(req: PdfToWordRequest) -> dict:
    """Convert a PDF to a Word (.docx) document. Handles both native text and scanned PDFs (via OCR)."""
    import fitz
    from docx import Document
    from docx.shared import Pt

    path = os.path.abspath(req.path)
    _check_safe(path)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail=f"File not found: {path}")

    output = os.path.abspath(req.output_path) if req.output_path else path.rsplit(".", 1)[0] + ".docx"
    _check_safe(output)
    os.makedirs(os.path.dirname(output), exist_ok=True)

    pdf = fitz.open(path)
    try:
        doc = Document()
        ocr_pages = 0

        for i, page in enumerate(pdf):
            if i > 0:
                doc.add_page_break()

            # Try native text extraction first
            blocks = page.get_text("dict")["blocks"]
            text_blocks = [b for b in blocks if b["type"] == 0]
            page_text = "".join(
                span["text"]
                for b in text_blocks
                for line in b["lines"]
                for span in line["spans"]
            ).strip()

            if page_text and len(page_text) > 20:
                # Native text — preserve formatting
                for block in text_blocks:
                    for line in block["lines"]:
                        line_text = "".join(span["text"] for span in line["spans"])
                        if not line_text.strip():
                            continue
                        para = doc.add_paragraph()
                        for span in line["spans"]:
                            run = para.add_run(span["text"])
                            run.font.size = Pt(span["size"])
                            if span["flags"] & 2 ** 0:
                                run.font.superscript = True
                            if span["flags"] & 2 ** 1:
                                run.font.italic = True
                            if span["flags"] & 2 ** 4:
                                run.font.bold = True
            else:
                # Scanned page — render to image and OCR
                ocr_pages += 1
                ocr_text = _ocr_pdf_page(page)
                if ocr_text:
                    for line in ocr_text.split("\n"):
                        if line.strip():
                            doc.add_paragraph(line)
                else:
                    doc.add_paragraph(f"[Page {i + 1}: OCR could not extract text]")

        total_pages = len(pdf)
    finally:
        pdf.close()
    doc.save(output)
    result = {"success": True, "path": output, "pages_converted": total_pages}
    if ocr_pages:
        result["ocr_pages"] = ocr_pages
        result["_hint"] = f"{ocr_pages} scanned page(s) converted via OCR."
    return result


def _ocr_pdf_page(page) -> str:
    """OCR a single PyMuPDF page by rendering to image and running OCR."""
    try:
        import fitz
        from PIL import Image

        from extractors import _HAS_TESSERACT, _get_gemini, _ocr_gemini, _ocr_tesseract, _preprocess_image

        # Render page at 200 DPI
        mat = fitz.Matrix(200 / 72, 200 / 72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img = _preprocess_image(img)

        gemini = _get_gemini()
        if gemini:
            return _ocr_gemini([img])
        elif _HAS_TESSERACT:
            return _ocr_tesseract([img])
        return ""
    except Exception:
        return ""
